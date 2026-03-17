#!/usr/bin/env python3
"""
Generate DA-branded Word document from Website Redesign Design Doc.
Output: reports/DA_Marketing_Website_Redesign_Design_Doc_2026-03-16.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand colors
DA_ORANGE = RGBColor(0xF3, 0x8B, 0x1C)
DA_BLACK = RGBColor(0x00, 0x00, 0x00)
DA_BLUE_GRAY = RGBColor(0x5A, 0x6B, 0x7A)
DA_GRAY = RGBColor(0xCB, 0xCB, 0xCB)
DA_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DA_LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)

FONT_NAME = "Arial"

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_PATH = os.path.join(BASE_DIR, "reports", "DA_Marketing_Website_Redesign_Design_Doc_2026-03-16.docx")


def set_run_font(run, size=11, bold=False, color=DA_BLACK, font_name=FONT_NAME):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)


def add_orange_rule(doc):
    """Add a thin orange horizontal rule."""
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="F8901E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_section_header(doc, text):
    """H1: DA Orange bold 16pt with orange rule below."""
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, color=DA_ORANGE)
    add_orange_rule(doc)


def add_subsection_header(doc, text):
    """H2: Blue-Gray bold 13pt."""
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=DA_BLUE_GRAY)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_bullet(doc, text, indent_level=0):
    """Add a bullet point, handling bold markdown markers."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
    p.space_after = Pt(2)
    _add_formatted_runs(p, text, size=11)
    return p


def add_checkbox_bullet(doc, text, indent_level=0):
    """Add a checkbox item using unicode ballot box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5 + indent_level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.space_after = Pt(2)
    run = p.add_run("\u2610 ")  # Unicode ballot box
    set_run_font(run, size=11)
    _add_formatted_runs(p, text, size=11)
    return p


def _add_formatted_runs(paragraph, text, size=11):
    """Parse **bold** markers and add runs accordingly."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        is_bold = (i % 2 == 1)
        set_run_font(run, size=size, bold=is_bold)


def add_table(doc, headers, rows):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=DA_WHITE)
        # Orange background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8901E" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_formatted_runs(p, val, size=10)
            # Alternating row shading
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    # Set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CBCBCB"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CBCBCB"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBCBCB"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CBCBCB"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CBCBCB"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CBCBCB"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # spacing after table
    return table


def add_footer(doc, text):
    """Add footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=8, color=DA_BLUE_GRAY)

        # Page number on the right — add tab stop then page field
        run2 = p.add_run("\t")
        set_run_font(run2, size=8, color=DA_BLUE_GRAY)

        # Right-align tab stop
        pPr = p._element.get_or_add_pPr()
        tabs = parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'  <w:tab w:val="right" w:pos="9360"/>'
            f'</w:tabs>'
        )
        pPr.append(tabs)

        # Page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run3 = p.add_run()
        set_run_font(run3, size=8, color=DA_BLUE_GRAY)
        run3._element.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run4 = p.add_run()
        set_run_font(run4, size=8, color=DA_BLUE_GRAY)
        run4._element.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5 = p.add_run()
        set_run_font(run5, size=8, color=DA_BLUE_GRAY)
        run5._element.append(fldChar2)


def build_document():
    doc = Document()

    # Page setup: letter, 0.75" margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = DA_BLACK

    # ── TITLE ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.space_after = Pt(4)
    run = title_p.add_run("Digital Accomplice Website Redesign — Design Doc")
    set_run_font(run, size=20, bold=True, color=DA_ORANGE)

    # Metadata line
    meta_p = doc.add_paragraph()
    meta_p.space_after = Pt(2)
    for label, value in [("Created: ", "2026-03-16"), ("  |  Status: ", "Planning phase — not yet building"), ("  |  Platform: ", "Webflow (new build) | Wix (current live site, kept as rollback)")]:
        run = meta_p.add_run(label)
        set_run_font(run, size=10, bold=True, color=DA_BLUE_GRAY)
        run = meta_p.add_run(value)
        set_run_font(run, size=10, color=DA_BLACK)

    add_orange_rule(doc)

    # ════════════════════════════════════════════
    # SECTION: What We Know
    # ════════════════════════════════════════════
    add_section_header(doc, "What We Know")

    add_subsection_header(doc, "Current State")
    for b in [
        "Live site: digitalaccomplice.com on Wix",
        "Site doesn't reflect new positioning (strategy + production, not just production)",
        "Need site that sells the 3-tier service model, showcases case studies, and leads with AI visibility angle",
    ]:
        add_bullet(doc, b)

    add_subsection_header(doc, "Platform Decision: Webflow")
    add_bullet(doc, "**Why Webflow over Wix:** Cleaner code (better SEO + AI visibility), Claude connector for assisted building, better CMS for blog/resources, more design control, zero maintenance")
    add_bullet(doc, "**Rollback plan:** Keep Wix site 100% intact and live. Build entire new site on Webflow staging URL. When ready, DNS swap. If problems, swap DNS back to Wix. Keep Wix active for 30 days post-launch.")
    add_bullet(doc, "**Claude connector:** Webflow has a Claude connector that may let us build/edit pages programmatically")

    add_subsection_header(doc, "Sitemap (5 pages + eyebrow nav)")
    add_body(doc, "Main nav: Home | About | Services | Work | Resources")
    add_body(doc, "Eyebrow nav: Contact | Privacy | Careers | Sitemap")

    # Page descriptions
    pages = [
        ("Page 1: Home", [
            "Hero: Positioning statement + hero video",
            "StoryBrand-style layout (see wireframe template in sitemap PDF)",
            "Problem \u2192 solution \u2192 CTA flow",
            "Three positioning statement options drafted (see sitemap PDF page 9, Simon\u2019s feedback)",
            "Multiple CTAs: free assessment, AI Visibility Snapshot, video interview, start a project",
        ]),
        ("Page 2: About", [
            "\"Why us\" section with Venn diagram (marketing agency + video production = DA)",
            "Origin story (Dane\u2019s background: NatGeo, Discovery, USATODAY, Xbox, PlayStation \u2192 DA in 2010)",
            "\"How we work\" \u2014 4-step proprietary process",
            "Team section (Dane + assembled experts)",
            "Client logos: Adobe, Twitch, Google, Microsoft, Honda, VW, Taco Bell, Samsung",
            "Dane\u2019s face on page for trust building",
        ]),
        ("Page 3: Services", [
            "3 tiers: Strategy ($5K/$10K/$15K), Pipeline (2-8 weeks), Project-Based (varies)",
            "Challenge/approach comparison table",
            "\"Three Revolutions\" framework (AI, Video, Business Culture)",
            "Sub-nav: Our Video Strategy Process, Pre-purposing Pipeline Setup, Project-Based Production",
            "Each service: What it is, Who it\u2019s for, How it works, What you get, Price/timing",
            "Budget allocation examples (\"$10K can be spent 3 different ways\")",
            "Note: Strategy comes first \u2014 \"you wouldn\u2019t build a house without a blueprint\"",
        ]),
        ("Page 4: Work", [
            "Logo wall at top (~10 B2B logos: AWS, Twitch, Google, Adobe)",
            "Three case study categories:",
            "  A) Video Strategy: Hinge, Thiel & Team, Crux",
            "  B) Video Pipeline: Twitch (sales enablement), Tube Mogul (outsourced team), Quatrain (field/remote)",
            "  C) One-off Content: AWS Cyber, Colton animated, Twitch, social clips",
            "Each pipeline case study needs workflow diagram graphic",
            "\"More examples\" button at bottom of one-offs",
        ]),
        ("Page 5: Resources", [
            "Blog (repurpose Substack newsletter as weekly posts + standalone blog posts)",
            "Newsletter subscribe/view (Substack link)",
            "Hourly consulting (hidden page)",
            "Webinars/events (link inactive until launched)",
            "Publications (3 Revolutions white paper, exec guide, quarterly cadence)",
            "Videos: thought leadership content archive",
        ]),
    ]

    for page_title, bullets in pages:
        p = doc.add_paragraph()
        p.space_before = Pt(8)
        p.space_after = Pt(2)
        run = p.add_run(page_title)
        set_run_font(run, size=11, bold=True, color=DA_BLACK)
        for b in bullets:
            indent = 1 if b.startswith("  ") else 0
            add_bullet(doc, b.strip(), indent_level=indent)

    add_subsection_header(doc, "CTA Strategy")
    add_bullet(doc, "**Primary:** Free assessment / book a call")
    add_bullet(doc, "**Secondary on-ramps** (lower friction):")
    add_bullet(doc, "AI Visibility Snapshot (lead capture \u2014 give email, get report)", indent_level=1)
    add_bullet(doc, "Video interview offer", indent_level=1)
    add_bullet(doc, "\"Start a video project\" inquiry", indent_level=1)
    add_bullet(doc, "CTAs woven throughout site, not just one page")
    add_bullet(doc, "No \"buy now\" button for strategy")

    add_subsection_header(doc, "Brand Guidelines")
    add_bullet(doc, "**Colors:** DA Orange #F8901E, Black #000, Blue-Gray #5A6B7A, Gray #CBCBCB, White #FFF, Light Gray #F5F5F5")
    add_bullet(doc, "**Fonts:** Inter/Arial/Helvetica")
    add_bullet(doc, "**Voice:** Direct, data-first, no-BS, short sentences")
    add_bullet(doc, "**Do NOT use:** #F5A623, #AAAAAA, #F0F0F0")

    add_subsection_header(doc, "Inspiration Sites")
    inspiration = [
        "storybrand.com \u2014 StoryBrand framework layout",
        "thursdaylabs.co \u2014 video prod company (\"I would buy from them\")",
        "rightfit.leadhunter.net",
        "flywheelos.com",
        "parakeeto.com",
        "chasingcreative.io \u2014 services section",
        "StoryBrand wireframe template (pribbledesign.com)",
        "promode.ai/the-use-cases/ \u2014 use cases layout idea",
    ]
    for site in inspiration:
        add_bullet(doc, site)

    # ════════════════════════════════════════════
    # SECTION: Open Questions
    # ════════════════════════════════════════════
    add_section_header(doc, "Open Questions")

    questions = [
        ["1", "Video hosting", "YouTube vs Wistia vs Vimeo vs hybrid for on-site video player? (Rec: hybrid \u2014 YouTube for discovery, Wistia/Vimeo for on-site embeds)"],
        ["2", "Blog approach", "Native Webflow blog CMS vs. embed/link to Substack? (Rec: Webflow native blog, cross-post to Substack)"],
        ["3", "Showcase page", "John Corcoran Showcase page idea \u2014 add as a service? Separate landing page?"],
        ["4", "Hero video", "What video goes in the hero? Does it exist yet or need to be produced?"],
        ["5", "Case study content", "Do we have the written case studies for Hinge, Thiel, Crux, Twitch, Tube Mogul, Quatrain? Or do these need to be written?"],
        ["6", "Workflow diagrams", "For pipeline case studies \u2014 who creates these? (Could be built as simple graphics in the site)"],
        ["7", "Team section", "Who else besides Dane? Need names, photos, bios?"],
        ["8", "Contact form", "What fields? Where does it submit to? (Calendly embed? Email?)"],
        ["9", "Analytics", "Google Analytics? Webflow native? Both?"],
        ["10", "Favicon/logo", "Current DA logo good to go, or updating?"],
        ["11", "Photos of Dane", "Professional headshot ready? Multiple photos for different sections?"],
        ["12", "Testimonials", "Any client quotes to include?"],
        ["13", "Publications", "Is the 3 Revolutions white paper written? Exec guide?"],
        ["14", "Hourly consulting", "What goes on the hidden page? Calendly link + rate?"],
    ]
    add_table(doc, ["#", "Topic", "Question / Notes"], questions)

    # ════════════════════════════════════════════
    # SECTION: Step-by-Step Build Plan
    # ════════════════════════════════════════════
    add_section_header(doc, "Step-by-Step Build Plan")

    phases = [
        ("Phase 0: Setup", [
            "0.1 \u2014 Dane creates Webflow account (free tier)",
            "0.2 \u2014 Create new blank Webflow project (\"Digital Accomplice\")",
            "0.3 \u2014 Set up Claude connector (follow Webflow blog post instructions)",
            "0.4 \u2014 Verify connector works (test a simple edit)",
            "0.5 \u2014 Set up global styles: DA colors, fonts, spacing system",
        ]),
        ("Phase 1: Home Page \u2014 First Draft", [
            "1.1 \u2014 Build page skeleton: nav, hero section, sections, footer",
            "1.2 \u2014 Nav bar: Home | About | Services | Work | Resources + eyebrow nav",
            "1.3 \u2014 Hero section: headline + subhead + CTA button (placeholder video)",
            "1.4 \u2014 Problem/solution section (StoryBrand style)",
            "1.5 \u2014 \"Three Revolutions\" overview section",
            "1.6 \u2014 Services preview (3 cards linking to Services page)",
            "1.7 \u2014 Logo wall (client logos)",
            "1.8 \u2014 CTA section (free assessment)",
            "1.9 \u2014 Footer (links, contact info, social)",
            "1.10 \u2014 **REVIEW WITH DANE** \u2014 adjust before moving on",
        ]),
        ("Phase 2: About Page", [
            "2.1 \u2014 \"Why us\" section + Venn diagram",
            "2.2 \u2014 Origin story (narrative copy)",
            "2.3 \u2014 \"How we work\" \u2014 4-step process (icons or numbered steps)",
            "2.4 \u2014 Team section (Dane bio + photo, team members TBD)",
            "2.5 \u2014 Full client logo wall",
            "2.6 \u2014 **REVIEW WITH DANE**",
        ]),
        ("Phase 3: Services Page", [
            "3.1 \u2014 \"Three ways we can work together\" overview",
            "3.2 \u2014 Strategy service detail (what/who/how/what you get/price)",
            "3.3 \u2014 Pipeline service detail",
            "3.4 \u2014 Project-based service detail",
            "3.5 \u2014 Challenge/approach comparison table",
            "3.6 \u2014 Budget allocation examples section",
            "3.7 \u2014 CTA (free assessment / book a call)",
            "3.8 \u2014 **REVIEW WITH DANE**",
        ]),
        ("Phase 4: Work Page", [
            "4.1 \u2014 Logo wall header",
            "4.2 \u2014 Case study category navigation",
            "4.3 \u2014 Strategy case studies (Hinge, Thiel, Crux) \u2014 need content from Dane",
            "4.4 \u2014 Pipeline case studies (Twitch, Tube Mogul, Quatrain) \u2014 need content + workflow diagrams",
            "4.5 \u2014 One-off showcase (AWS Cyber, Colton, Twitch, social clips) \u2014 need video embeds",
            "4.6 \u2014 **REVIEW WITH DANE**",
        ]),
        ("Phase 5: Resources Page", [
            "5.1 \u2014 Blog setup (Webflow CMS collection)",
            "5.2 \u2014 Import/create first 3-5 blog posts from Substack newsletter",
            "5.3 \u2014 Newsletter subscribe section (Substack link or embedded form)",
            "5.4 \u2014 Publications section (white paper, exec guide \u2014 PDFs or landing pages)",
            "5.5 \u2014 Videos section (thought leadership archive)",
            "5.6 \u2014 Webinars/events placeholder (inactive link)",
            "5.7 \u2014 Hourly consulting hidden page",
            "5.8 \u2014 **REVIEW WITH DANE**",
        ]),
        ("Phase 6: Secondary Pages + Polish", [
            "6.1 \u2014 Contact page (form + Calendly embed)",
            "6.2 \u2014 Privacy policy page",
            "6.3 \u2014 Careers page (placeholder or real)",
            "6.4 \u2014 Sitemap page",
            "6.5 \u2014 404 page",
            "6.6 \u2014 Mobile responsive pass (test all pages on phone/tablet)",
            "6.7 \u2014 **FULL REVIEW WITH DANE**",
        ]),
        ("Phase 7: Video Hosting Decision + Implementation", [
            "7.1 \u2014 Decide on video hosting (YouTube/Wistia/Vimeo/hybrid)",
            "7.2 \u2014 Set up chosen platform account",
            "7.3 \u2014 Upload hero video",
            "7.4 \u2014 Upload case study videos",
            "7.5 \u2014 Upload thought leadership videos",
            "7.6 \u2014 Replace all placeholder video embeds with real ones",
            "7.7 \u2014 Test video playback on desktop + mobile",
        ]),
        ("Phase 8: SEO + AI Visibility", [
            "8.1 \u2014 Page titles + meta descriptions for all pages",
            "8.2 \u2014 Open Graph tags (social sharing previews)",
            "8.3 \u2014 Schema markup (Organization, Service, Video)",
            "8.4 \u2014 Alt text for all images",
            "8.5 \u2014 Sitemap.xml generation",
            "8.6 \u2014 Google Analytics / Search Console setup",
            "8.7 \u2014 Test AI visibility (ask ChatGPT/Perplexity about DA \u2014 baseline before launch)",
        ]),
        ("Phase 9: Launch Prep", [
            "9.1 \u2014 Final content review \u2014 all copy proofread",
            "9.2 \u2014 Test all forms (contact, newsletter, CTAs)",
            "9.3 \u2014 Test all links (internal + external)",
            "9.4 \u2014 Speed test (PageSpeed Insights)",
            "9.5 \u2014 Cross-browser test (Chrome, Safari, Firefox)",
            "9.6 \u2014 Document rollback procedure (DNS swap steps)",
            "9.7 \u2014 **FINAL SIGN-OFF FROM DANE**",
        ]),
        ("Phase 10: Go Live", [
            "10.1 \u2014 Point digitalaccomplice.com DNS to Webflow",
            "10.2 \u2014 Verify site is live and working",
            "10.3 \u2014 Test all forms again on live domain",
            "10.4 \u2014 Submit sitemap to Google Search Console",
            "10.5 \u2014 Monitor for 48 hours",
            "10.6 \u2014 If all good, keep Wix as backup for 30 more days, then cancel",
        ]),
    ]

    for phase_title, steps in phases:
        add_subsection_header(doc, phase_title)
        for step in steps:
            add_checkbox_bullet(doc, step)

    # ════════════════════════════════════════════
    # SECTION: Rules for This Build
    # ════════════════════════════════════════════
    add_section_header(doc, "Rules for This Build")

    rules = [
        ["1", "Small steps", "Never build more than one section without reviewing with Dane."],
        ["2", "No rushing", "Quality over speed. Dane said \"ASAP but don't railroad it.\""],
        ["3", "Review at every phase", "Each phase ends with a Dane review checkpoint."],
        ["4", "Rollback always available", "Wix stays live until Webflow is fully tested and approved."],
        ["5", "Content first, design second", "Get the words right, then make it look good."],
        ["6", "Video hosting decided in Phase 7", "Don't let that decision block progress."],
        ["7", "Each session = 1-2 steps max", "Avoid context drift. Save, commit, fresh start."],
    ]
    add_table(doc, ["#", "Rule", "Detail"], rules)

    # ── FOOTER ──
    add_footer(doc, "Digital Accomplice | Website Redesign Design Doc | March 2026")

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
